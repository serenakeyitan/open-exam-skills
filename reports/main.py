"""
Reports Skill - Generate structured reports from research
"""

import json
import argparse
import os
from typing import Optional
from pydantic import BaseModel
from loguru import logger
import sys

logger.remove()
logger.add(sys.stderr, level="INFO")


class Report(BaseModel):
    """Report structure."""
    title: str
    executive_summary: str
    sections: list[dict]
    conclusions: str
    recommendations: list[str]


def generate_with_ai(content: str, report_type: str, api_key: str, provider: str = "gemini") -> str:
    """Generate report using AI."""
    prompt = f"""Create a professional {report_type} report from this content.

Content: {content[:15000]}

Generate a structured report with:
- title: Report title
- executive_summary: Brief overview (2-3 paragraphs)
- sections: Array of {{title, content}} objects
- conclusions: Key findings and insights
- recommendations: Action items

Return as JSON:
{{
  "title": "Report Title",
  "executive_summary": "Summary text...",
  "sections": [
    {{"title": "Section 1", "content": "Section content..."}},
    {{"title": "Section 2", "content": "Section content..."}}
  ],
  "conclusions": "Conclusions text...",
  "recommendations": ["Rec 1", "Rec 2"]
}}"""

    if provider == "gemini":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-3-pro-preview")
        response = model.generate_content(prompt)
        return response.text
    else:
        if not args.input:
            parser.error("--input is required when not in test mode")
        from anthropic import Anthropic
        client = Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text


def export_to_markdown(report: Report, output_path: str) -> None:
    """Export report as Markdown."""
    md = f"# {report.title}\n\n"
    md += f"## Executive Summary\n\n{report.executive_summary}\n\n"

    for section in report.sections:
        md += f"## {section['title']}\n\n{section['content']}\n\n"

    md += f"## Conclusions\n\n{report.conclusions}\n\n"

    if report.recommendations:
        md += "## Recommendations\n\n"
        for i, rec in enumerate(report.recommendations, 1):
            md += f"{i}. {rec}\n"

    with open(output_path, 'w') as f:
        f.write(md)

    logger.info(f"Markdown report saved: {output_path}")


def export_to_pdf(report: Report, output_path: str) -> None:
    """Export report as PDF."""
    try:
        import markdown2
        from weasyprint import HTML

        # Generate markdown first
        md_path = output_path.replace('.pdf', '.md')
        export_to_markdown(report, md_path)

        # Convert to HTML
        with open(md_path) as f:
            md_content = f.read()

        html_content = markdown2.markdown(md_content)
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: Arial; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
    </style>
</head>
<body>{html_content}</body>
</html>
"""

        # Convert to PDF
        HTML(string=html).write_pdf(output_path)
        logger.info(f"PDF report saved: {output_path}")

    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        logger.info("Falling back to Markdown")
        md_path = output_path.replace('.pdf', '.md')
        export_to_markdown(report, md_path)


def export_to_docx(report: Report, output_path: str) -> None:
    """Export report as DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading(report.title, 0)
    title.runs[0].font.size = Pt(24)

    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(report.executive_summary)

    # Sections
    for section in report.sections:
        doc.add_heading(section['title'], 1)
        doc.add_paragraph(section['content'])

    # Conclusions
    doc.add_heading('Conclusions', 1)
    doc.add_paragraph(report.conclusions)

    # Recommendations
    if report.recommendations:
        doc.add_heading('Recommendations', 1)
        for rec in report.recommendations:
            doc.add_paragraph(rec, style='List Number')

    doc.save(output_path)
    logger.info(f"DOCX report saved: {output_path}")


def generate_report(
    content: str,
    report_type: str = "executive_summary",
    output_path: str = "report.md"
) -> str:
    """Generate report from content."""
    logger.info(f"Generating {report_type} report...")

    # Get API key
    gemini_key = os.getenv("GEMINI_API_KEY", "")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")

    if gemini_key:
        response = generate_with_ai(content, report_type, gemini_key, "gemini")
    elif anthropic_key:
        response = generate_with_ai(content, report_type, anthropic_key, "anthropic")
    else:
        if not args.input:
            parser.error("--input is required when not in test mode")
        raise ValueError("No API key found")

    # Parse response
    text = response.strip()
    if text.startswith("```json"):
        text = text[7:]
    if text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]

    data = json.loads(text.strip())
    report = Report(**data)

    # Export based on file extension
    if output_path.endswith('.pdf'):
        export_to_pdf(report, output_path)
    elif output_path.endswith('.docx'):
        export_to_docx(report, output_path)
    else:
        export_to_markdown(report, output_path)

    return output_path


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", "-i", required=True)
    parser.add_argument("--output", "-o", default="report.md")
    parser.add_argument("--type", "-t", default="executive_summary")
    parser.add_argument("--test", action="store_true")

    args = parser.parse_args()

    if args.test:
        content = "Quantum computing represents a paradigm shift. Key findings: superposition enables parallel computation, entanglement provides quantum correlation, error correction remains challenging."
        result = generate_report(content, "research_brief", "test_report.md")
        print(f"Test report: {result}")
    else:
        if not args.input:
            parser.error("--input is required when not in test mode")
        with open(args.input) as f:
            content = f.read()
        result = generate_report(content, args.type, args.output)
        print(f"Report generated: {result}")


if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv()
    main()
